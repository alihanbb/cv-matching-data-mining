"""Import sister-cloned public repos into Bronze JSONL (import-time only).

External layouts are parsed here; runtime pipeline code reads only Bronze JSONL.
"""

from __future__ import annotations

import argparse
import ast
import csv
import json
import logging
import re
import sys
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass
from json import JSONDecodeError
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

REPO_ALIASES: dict[str, list[str]] = {
    "vanetik": ["vacancy-resume-matching-dataset"],
    "dataturks": ["Entity-Recognition-In-Resumes-SpaCy"],
    "mehyar": ["NER-Annotated-CVs"],
    "nlp_ner": ["NLP_NER_ON_RESUME"],
}

SOURCE_TAGS: dict[str, tuple[str, ...]] = {
    "vanetik": ("vacancy_resume_matching",),
    "dataturks": ("dataturks_resume_ner_train", "dataturks_resume_ner_test"),
    "mehyar": ("mehyar_ner_annotated_cv",),
    "nlp_ner": ("nlp_ner_on_resume_json_demo",),
}


def resolve_repo(source_root: Path, key: str) -> Path | None:
    for name in REPO_ALIASES.get(key, ()):
        p = (source_root / name).resolve()
        if p.is_dir():
            return p
    return None


def _safe_str(v: Any) -> str:
    return "" if v is None else str(v).strip()


def stats_for_records(
    records: Iterable[dict[str, Any]], source_key: str = "source"
) -> dict[str, Any]:
    by_source: Counter[str] = Counter()
    for rec in records:
        by_source[_safe_str(rec.get(source_key))] += 1
    return {"total": sum(by_source.values()), "by_source": dict(by_source)}


def write_stats(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def load_jsonl_by_id(path: Path, id_key: str) -> dict[str, dict[str, Any]]:
    out: dict[str, dict[str, Any]] = {}
    if not path.is_file():
        return out
    with path.open(encoding="utf-8", errors="replace", newline="\n") as f:
        for ln, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except JSONDecodeError:
                logger.warning("Skipping malformed JSONL line %s:%d", path, ln)
                continue
            if isinstance(obj, dict):
                rid = _safe_str(obj.get(id_key))
                if rid:
                    out[rid] = obj
    return out


def filter_out_sources(
    rows: dict[str, dict[str, Any]], tags: Iterable[str]
) -> dict[str, dict[str, Any]]:
    tag_set = {t.strip() for t in tags if t}
    return {k: v for k, v in rows.items() if _safe_str(v.get("source")) not in tag_set}


def _utf8_json_safe(obj: Any) -> Any:
    """Recursively replace lone UTF-16 surrogates so ``json`` + UTF-8 file write cannot fail."""
    if isinstance(obj, dict):
        return {k: _utf8_json_safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_utf8_json_safe(v) for v in obj]
    if isinstance(obj, str):
        return "".join("\ufffd" if 0xD800 <= ord(ch) <= 0xDFFF else ch for ch in obj)
    return obj


def write_jsonl(path: Path, records: Iterable[dict[str, Any]]) -> int:
    path.parent.mkdir(parents=True, exist_ok=True)
    n = 0
    with path.open("w", encoding="utf-8", newline="\n") as f:
        for rec in records:
            safe = _utf8_json_safe(rec)
            f.write(json.dumps(safe, ensure_ascii=False) + "\n")
            n += 1
    return n


def _extract_bracket_value(text: str, var: str) -> str | None:
    """Return substring ``[[...], ...]`` assigned to ``var`` (balanced brackets)."""
    needle = f"{var}="
    i = text.find(needle)
    if i < 0:
        return None
    i += len(needle)
    start = text.find("[", i)
    if start < 0:
        return None
    depth = 0
    for j in range(start, len(text)):
        ch = text[j]
        if ch == "[":
            depth += 1
        elif ch == "]":
            depth -= 1
            if depth == 0:
                return text[start : j + 1]
    return None


def parse_vanetik_annotation_rankings(
    txt_path: Path,
) -> tuple[list[list[int]], list[list[int]]] | None:
    if not txt_path.is_file():
        return None
    raw = txt_path.read_text(encoding="utf-8", errors="replace")
    s1, s2 = _extract_bracket_value(raw, "ANNOTATOR_1_RANKINGS"), _extract_bracket_value(
        raw, "ANNOTATOR_2_RANKINGS"
    )
    if not s1 or not s2:
        logger.warning("Could not locate ANNOTATOR_*_RANKINGS lists in %s", txt_path.name)
        return None
    try:
        r1 = ast.literal_eval(s1)
        r2 = ast.literal_eval(s2)
    except (SyntaxError, ValueError) as exc:
        logger.warning("ast.literal_eval failed for rankings (%s): %s", txt_path.name, exc)
        return None
    if not (
        isinstance(r1, list)
        and isinstance(r2, list)
        and len(r1) == len(r2)
        and all(isinstance(x, list) and len(x) == 5 for x in r1)
        and all(isinstance(x, list) and len(x) == 5 for x in r2)
    ):
        logger.warning("Unexpected ranking structure in %s", txt_path.name)
        return None
    return r1, r2


def rank_value_to_relevance(r: float) -> int:
    """Lower rank value = better match (1 best on 1..5 scale)."""
    ir = int(round(r + 1e-6))
    if ir <= 1:
        return 3
    if ir == 2:
        return 2
    if ir == 3:
        return 1
    return 0


def ground_truth_from_rankings(
    rankings: tuple[list[list[int]], list[list[int]]],
    job_ids_ordered: list[str],
) -> list[dict[str, Any]]:
    a1, a2 = rankings
    rows: list[dict[str, Any]] = []
    for cv_i, (row1, row2) in enumerate(zip(a1, a2, strict=True)):
        rid = f"vanetik_cv_{cv_i + 1:03d}"
        for v_idx in range(5):
            r_mean = (row1[v_idx] + row2[v_idx]) / 2.0
            rel = rank_value_to_relevance(r_mean)
            rows.append(
                {
                    "job_id": job_ids_ordered[v_idx],
                    "resume_id": rid,
                    "relevance": rel,
                    "source": "vacancy_resume_matching",
                }
            )
    return rows


def _docx_to_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip()).strip()


def _vanetik_resume_id_from_name(name: str) -> str | None:
    digits = "".join(ch for ch in Path(name).stem if ch.isdigit())
    if not digits:
        return None
    return f"vanetik_cv_{int(digits):03d}"


def _vacancy_title_and_text(row: dict[str, str]) -> tuple[str, str]:
    lower = {k.lower().strip(): (v or "") for k, v in row.items() if k}
    title = ""
    for key in lower:
        if "job_title" in key or key == "title":
            title = _safe_str(lower[key])
            if title:
                break
    body = ""
    for key, val in lower.items():
        if any(x in key for x in ("description", "vacancy", "job_text", "text")) and len(val) > len(
            body
        ):
            body = _safe_str(val)
    if not body:
        for key, val in lower.items():
            if "job" in key and len(val) > len(body):
                body = _safe_str(val)
    raw = "\n\n".join(p for p in (title, body) if p).strip()
    return title, raw


def import_vanetik(repo: Path) -> tuple[list[dict], list[dict], list[dict]]:
    resumes: list[dict[str, Any]] = []
    jobs: list[dict[str, Any]] = []
    gt_template: list[dict[str, Any]] = []

    cv_dir = repo / "CV"
    if not cv_dir.is_dir():
        logger.warning("Vanetik CV folder missing: %s", cv_dir)
    else:
        for docx in sorted(cv_dir.glob("*.docx")):
            rid = _vanetik_resume_id_from_name(docx.name)
            if not rid:
                logger.warning("Skip DOCX (no numeric id in name): %s", docx.name)
                continue
            try:
                raw_text = _docx_to_text(docx)
            except Exception as exc:
                logger.warning("Skip DOCX (read failed) %s: %s", docx.name, exc)
                continue
            if not raw_text:
                logger.warning("Skip DOCX (empty text): %s", docx.name)
                continue
            resumes.append(
                {
                    "resume_id": rid,
                    "source": "vacancy_resume_matching",
                    "source_file": docx.name,
                    "raw_text": raw_text,
                    "language": "en",
                    "labels": {"entities": []},
                    "metadata": {"original_format": "docx", "category": None},
                }
            )

    vac_csv = repo / "5_vacancies.csv"
    job_ids_order: list[str] = []
    if not vac_csv.is_file():
        logger.warning("Vanetik vacancies CSV missing: %s", vac_csv)
    else:
        with vac_csv.open(encoding="utf-8-sig", newline="", errors="replace") as f:
            reader = csv.DictReader(f)
            for idx, row in enumerate(reader, start=1):
                jid = f"vanetik_vacancy_{idx:03d}"
                job_ids_order.append(jid)
                title, raw_text = _vacancy_title_and_text(row)
                if not raw_text:
                    logger.warning("Skip vacancy row %d (empty text)", idx)
                    continue
                jobs.append(
                    {
                        "job_id": jid,
                        "source": "vacancy_resume_matching",
                        "source_file": vac_csv.name,
                        "raw_text": raw_text,
                        "title": title or jid,
                        "language": "en",
                        "metadata": {
                            "original_format": "csv",
                            "uid": _safe_str(row.get("uid")),
                        },
                    }
                )

    ann = repo / "annotations-for-the-first-30-vacancies.txt"
    parsed = parse_vanetik_annotation_rankings(ann)
    gt_template: list[dict[str, Any]] = []
    if parsed and len(job_ids_order) >= 5:
        gt_template = ground_truth_from_rankings(parsed, job_ids_order[:5])
    elif not parsed and job_ids_order:
        for j in range(min(5, len(job_ids_order))):
            for i in range(30):
                gt_template.append(
                    {
                        "job_id": job_ids_order[j],
                        "resume_id": f"vanetik_cv_{i + 1:03d}",
                        "relevance": "",
                        "source": "vacancy_resume_matching",
                    }
                )

    return resumes, jobs, gt_template


def _dataturks_points_to_entities(content: str, annotation: Any) -> list[dict[str, Any]]:
    if not isinstance(annotation, list):
        return []
    out: list[dict[str, Any]] = []
    for ann in annotation:
        if not isinstance(ann, dict):
            continue
        pts = ann.get("points")
        if not isinstance(pts, list) or not pts:
            continue
        p0 = pts[0]
        if not isinstance(p0, dict):
            continue
        s_raw, e_raw = p0.get("start"), p0.get("end")
        if s_raw is None or e_raw is None:
            continue
        s_i, e_i = int(s_raw), int(e_raw)
        # DataTurks uses inclusive span end in many exports — convert to slice end exclusive:
        slice_end = e_i + 1 if e_i >= s_i else e_i
        tx = _ensure_entity_span(content, s_i, slice_end)

        labs = ann.get("label")
        lab_items: Iterable[Any]
        if isinstance(labs, list):
            lab_items = labs
        elif labs is None:
            continue
        else:
            lab_items = (labs,)
        for lab in lab_items:
            if lab is None:
                continue
            rec = {"start": s_i, "end": slice_end, "label": _safe_str(lab), "text": tx}
            out.append(rec)
    return out


def _ensure_entity_span(text: str, start: int, end: int) -> str:
    if start < 0 or end > len(text) or start > end:
        return ""
    return text[start:end]


def dataturks_iter(path: Path, source: str, split_tag: str) -> tuple[list[dict], list[dict]]:
    resumes: list[dict[str, Any]] = []
    anns: list[dict[str, Any]] = []
    if not path.is_file():
        logger.warning("DataTurks file missing: %s", path)
        return resumes, anns
    seq_i = 0
    with path.open(encoding="utf-8", errors="replace", newline="\n") as f:
        for idx, line in enumerate(f, start=1):
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
            except JSONDecodeError:
                logger.warning("Skip DataTurks line %d (JSON): %s", idx, path.name)
                continue
            if not isinstance(item, dict):
                continue
            content = _safe_str(item.get("content"))
            if not content:
                continue
            seq_i += 1
            prefix = "train" if "train" in split_tag else "test"
            rid = f"dataturks_{prefix}_{seq_i:06d}"
            aid = f"dataturks_{prefix}_ann_{seq_i:06d}"
            ents = _dataturks_points_to_entities(content, item.get("annotation"))
            norm_ents: list[dict[str, Any]] = []
            for e in ents:
                if not _safe_str(e.get("label")):
                    continue
                norm_ents.append(e)
            resumes.append(
                {
                    "resume_id": rid,
                    "source": source,
                    "source_file": path.name,
                    "raw_text": content,
                    "language": "en",
                    "labels": {"entities": norm_ents},
                    "metadata": {
                        "original_format": "json",
                        "split": prefix,
                        "category": None,
                    },
                }
            )
            anns.append(
                {
                    "annotation_id": aid,
                    "source": source,
                    "source_file": path.name,
                    "text": content,
                    "entities": norm_ents,
                    "metadata": {"original_format": "json", "split": prefix},
                }
            )
    return resumes, anns


def mehyar_find_json_root(repo: Path) -> Path | None:
    for rel in (
        Path("extracted") / "ResumesJsonAnnotated",
        Path("ResumesJsonAnnotated"),
    ):
        p = repo / rel
        if p.is_dir():
            return p
    return None


def _mehyar_triples_to_entities(text: str, raw: Any) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    if isinstance(raw, list):
        if raw and isinstance(raw[0], (list, tuple)) and len(raw[0]) >= 3:
            for t in raw:
                if isinstance(t, (list, tuple)) and len(t) >= 3:
                    s_i, e_i, lab = int(t[0]), int(t[1]), _safe_str(t[2])
                    if not lab:
                        continue
                    tx = _ensure_entity_span(text, s_i, e_i)
                    out.append({"start": s_i, "end": e_i, "label": lab, "text": tx})
            return out
        for d in raw:
            if not isinstance(d, dict):
                continue
            s_raw = d.get("start") or d.get("begin") or d.get("from")
            e_raw = d.get("end") or d.get("to")
            lab = _safe_str(d.get("label") or d.get("type") or d.get("entity"))
            if lab == "":
                continue
            if s_raw is None or e_raw is None:
                ent_text = _safe_str(d.get("text"))
                out.append({"start": None, "end": None, "label": lab, "text": ent_text})
                continue
            s_i, e_i = int(s_raw), int(e_raw)
            tx = d.get("text")
            tx = _safe_str(tx) if tx else _ensure_entity_span(text, s_i, e_i)
            out.append({"start": s_i, "end": e_i, "label": lab, "text": tx})
    return out


def import_mehyar(repo: Path) -> tuple[list[dict], list[dict]]:
    resumes: list[dict[str, Any]] = []
    anns: list[dict[str, Any]] = []
    root = mehyar_find_json_root(repo)
    if root is None:
        logger.warning("Mehyar ResumesJsonAnnotated directory not found under %s", repo)
        return resumes, anns

    sorted_paths = sorted(root.glob("*.json"))
    for idx, path in enumerate(sorted_paths, start=1):
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        except (JSONDecodeError, OSError) as exc:
            logger.warning("Skip Mehyar JSON %s: %s", path.name, exc)
            continue
        if not isinstance(data, dict):
            continue
        text = _safe_str(data.get("text"))
        if not text:
            continue
        rid = f"mehyar_{idx:06d}"
        aid = f"mehyar_ann_{idx:06d}"
        raw_ann = data.get("annotations")
        ents = _mehyar_triples_to_entities(text, raw_ann)
        resumes.append(
            {
                "resume_id": rid,
                "source": "mehyar_ner_annotated_cv",
                "source_file": path.name,
                "raw_text": text,
                "language": "en",
                "labels": {"entities": ents},
                "metadata": {"original_format": "json", "category": None},
            }
        )
        anns.append(
            {
                "annotation_id": aid,
                "source": "mehyar_ner_annotated_cv",
                "source_file": path.name,
                "text": text,
                "entities": ents,
                "metadata": {"original_format": "json"},
            }
        )
    return resumes, anns


def json_resume_extended_text(obj: dict[str, Any]) -> tuple[str, dict[str, Any]]:
    """Flatten JSON Resume-ish dict; return (plain_text, structured_metadata_shell)."""

    meta: dict[str, Any] = {"structured": {}, "schema": "json_resume_like"}
    parts: list[str] = []

    def add_heading(title: str, body: str) -> None:
        if body.strip():
            parts.append(f"{title}\n{body}".strip())

    basics = obj.get("basics") if isinstance(obj.get("basics"), dict) else {}
    if basics:
        meta["structured"]["basics"] = basics
        for key in ("name", "label", "email", "phone", "summary"):
            v = basics.get(key)
            if v and _safe_str(v):
                parts.append(_safe_str(v))
        loc = basics.get("location")
        if isinstance(loc, dict) and loc.get("address"):
            parts.append(_safe_str(loc["address"]))

    for section_key in ("education",):
        seq = obj.get(section_key)
        if not isinstance(seq, list):
            continue
        meta["structured"][section_key] = seq
        for edu in seq:
            if not isinstance(edu, dict):
                continue
            chunk = " ".join(
                _safe_str(edu.get(k))
                for k in ("studyType", "area", "institution", "score")
                if edu.get(k)
            )
            if chunk:
                parts.append(chunk)

    seq = obj.get("work") or []
    meta["structured"]["work"] = seq
    for work in seq:
        if not isinstance(work, dict):
            continue
        header = " ".join(_safe_str(work.get(k)) for k in ("position", "name") if work.get(k))
        if header:
            parts.append(header)
        for hl in work.get("highlights") or []:
            if _safe_str(hl):
                parts.append(_safe_str(hl))
        u = work.get("unlabeled")
        if _safe_str(u):
            parts.append(_safe_str(u))

    seq = obj.get("skills") or []
    meta["structured"]["skills"] = seq
    for sk in seq:
        if not isinstance(sk, dict):
            continue
        kws = sk.get("keywords")
        if isinstance(kws, list) and kws:
            parts.append(", ".join(_safe_str(x) for x in kws if x))
        nm = sk.get("name")
        if _safe_str(nm):
            parts.append(_safe_str(nm))

    seq = obj.get("projects") or []
    meta["structured"]["projects"] = seq
    for pr in seq:
        if not isinstance(pr, dict):
            continue
        hdr = _safe_str(pr.get("name") or pr.get("title"))
        if hdr:
            parts.append(hdr)
        desc = _safe_str(pr.get("description") or pr.get("summary"))
        if desc:
            parts.append(desc)
        for hl in pr.get("highlights") or []:
            if _safe_str(hl):
                parts.append(_safe_str(hl))

    seq = obj.get("certificates") or []
    meta["structured"]["certificates"] = seq
    for c in seq:
        if isinstance(c, dict):
            parts.append(_safe_str(c.get("name") or c.get("issuer")))
        elif _safe_str(c):
            parts.append(_safe_str(c))

    seq = obj.get("languages") or []
    meta["structured"]["languages"] = seq
    for lang in seq:
        if isinstance(lang, dict):
            lab = _safe_str(lang.get("language"))
            flu = _safe_str(lang.get("fluency"))
            parts.append(f"{lab} {flu}".strip() if lab else flu)
        elif _safe_str(lang):
            parts.append(_safe_str(lang))

    text = "\n\n".join(p for p in parts if p).strip()
    return text, meta


def import_nlp_resume_json(path: Path) -> list[dict[str, Any]]:
    if not path.is_file():
        logger.warning("NLP_NER resume.json missing: %s", path)
        return []
    try:
        obj = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except (JSONDecodeError, OSError) as exc:
        logger.warning("Could not read resume.json: %s", exc)
        return []
    if not isinstance(obj, dict):
        return []
    text, meta = json_resume_extended_text(obj)
    if not text:
        logger.warning("NLP resume.json produced empty text")
        return []
    return [
        {
            "resume_id": "nlp_ner_resume_001",
            "source": "nlp_ner_on_resume_json_demo",
            "source_file": path.name,
            "raw_text": text,
            "language": "en",
            "labels": {"entities": []},
            "metadata": {"original_format": "json", **meta},
        }
    ]


@dataclass
class ImportOutputs:
    project_root: Path
    resumes_path: Path
    jobs_path: Path
    ner_path: Path
    resumes_stats: Path
    jobs_stats: Path
    ner_stats: Path
    eval_dir: Path
    gt_csv: Path
    gt_tpl: Path
    gt_guide_hint: Path


def default_outputs(root: Path) -> ImportOutputs:
    bronze = root / "data" / "bronze"
    return ImportOutputs(
        project_root=root,
        resumes_path=bronze / "resumes" / "resumes_bronze.jsonl",
        jobs_path=bronze / "jobs" / "jobs_bronze.jsonl",
        ner_path=bronze / "annotations" / "ner_annotations_bronze.jsonl",
        resumes_stats=bronze / "resumes" / "resumes_bronze.stats.json",
        jobs_stats=bronze / "jobs" / "jobs_bronze.stats.json",
        ner_stats=bronze / "annotations" / "ner_annotations_bronze.stats.json",
        eval_dir=root / "data" / "evaluation",
        gt_csv=root / "data" / "evaluation" / "ground_truth.csv",
        gt_tpl=root / "data" / "evaluation" / "ground_truth_template.csv",
        gt_guide_hint=root / "docs" / "GROUND_TRUTH_GUIDE.md",
    )


def run_import(
    *,
    source_root: Path,
    project_root: Path,
    keys: Iterable[str],
    overwrite: bool,
) -> dict[str, Any]:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    outs = default_outputs(project_root)
    outs.eval_dir.mkdir(parents=True, exist_ok=True)

    key_list = sorted(set(keys))

    resumes = (
        load_jsonl_by_id(outs.resumes_path, "resume_id") if outs.resumes_path.is_file() else {}
    )
    jobs = load_jsonl_by_id(outs.jobs_path, "job_id")
    anns = load_jsonl_by_id(outs.ner_path, "annotation_id")

    all_keys_set = set(REPO_ALIASES.keys())
    if overwrite:
        if set(key_list) == all_keys_set:
            resumes.clear()
            jobs.clear()
            anns.clear()
        else:
            strip_tags = {t for k in key_list for t in SOURCE_TAGS.get(k, ())}
            resumes = {
                rk: rv
                for rk, rv in resumes.items()
                if _safe_str(rv.get("source")) not in strip_tags
            }
            jobs = {
                jk: jv for jk, jv in jobs.items() if _safe_str(jv.get("source")) not in strip_tags
            }
            anns = {
                ak: av for ak, av in anns.items() if _safe_str(av.get("source")) not in strip_tags
            }

    gt_rows_written: list[dict[str, Any]] = []
    had_auto_gt = False

    for key in sorted(key_list):
        repo = resolve_repo(source_root, key)
        if repo is None:
            logger.warning(
                "Skipping %s — sibling repo folder not found under %s.",
                key,
                source_root,
            )
            continue

        if key == "vanetik":
            rlist, jlist, gt_maybe = import_vanetik(repo)
            for r in rlist:
                resumes[r["resume_id"]] = r
            for j in jlist:
                jobs[j["job_id"]] = j
            if gt_maybe and isinstance(gt_maybe[0].get("relevance"), int):
                gt_rows_written = gt_maybe
                had_auto_gt = True
                outs.gt_tpl.unlink(missing_ok=True)
            elif gt_maybe:
                if overwrite or not outs.gt_tpl.is_file():
                    write_gt_template_csv(outs.gt_tpl, gt_maybe)

        elif key == "dataturks":
            tr_r, tr_a = dataturks_iter(
                repo / "traindata.json", "dataturks_resume_ner_train", "train"
            )
            te_r, te_a = dataturks_iter(repo / "testdata.json", "dataturks_resume_ner_test", "test")
            for r in tr_r + te_r:
                resumes[r["resume_id"]] = r
            for a in tr_a + te_a:
                anns[a["annotation_id"]] = a

        elif key == "mehyar":
            mr, ma = import_mehyar(repo)
            for r in mr:
                resumes[r["resume_id"]] = r
            for a in ma:
                anns[a["annotation_id"]] = a

        elif key == "nlp_ner":
            for r in import_nlp_resume_json(repo / "resume.json"):
                resumes[r["resume_id"]] = r

    write_jsonl(outs.resumes_path, sorted(resumes.values(), key=lambda x: x["resume_id"]))
    write_jsonl(outs.jobs_path, sorted(jobs.values(), key=lambda x: x["job_id"]))
    write_jsonl(outs.ner_path, sorted(anns.values(), key=lambda x: x["annotation_id"]))

    write_stats(outs.resumes_stats, stats_for_records(resumes.values()))
    write_stats(outs.jobs_stats, stats_for_records(jobs.values()))
    write_stats(outs.ner_stats, stats_for_records(anns.values()))

    if gt_rows_written and had_auto_gt:
        with outs.gt_csv.open("w", encoding="utf-8", newline="\n") as f:
            w = csv.DictWriter(f, fieldnames=["job_id", "resume_id", "relevance", "source"])
            w.writeheader()
            for row in gt_rows_written:
                w.writerow(row)

    return {
        "resumes": len(resumes),
        "jobs": len(jobs),
        "annotations": len(anns),
        "ground_truth_rows": len(gt_rows_written),
    }


def write_gt_template_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as f:
        w = csv.DictWriter(f, fieldnames=["job_id", "resume_id", "relevance", "source"])
        w.writeheader()
        for row in rows:
            w.writerow(row)


def run_cli(argv: list[str] | None = None, *, default_project_root: Path | None = None) -> None:
    ap = argparse.ArgumentParser(description="Import sister-cloned CV/NER repos into Bronze JSONL.")
    ap.add_argument(
        "--source-root",
        type=Path,
        default=Path(".."),
        help="Directory that contains the four sibling repo folders (default: parent of CWD).",
    )
    ap.add_argument(
        "--project-root",
        type=Path,
        default=None,
        help="cv-matching-data-mining root (default: parent of this script if launched from scripts/).",
    )
    ap.add_argument("--all", action="store_true", help="Import every supported source.")
    ap.add_argument(
        "--source",
        action="append",
        choices=sorted(REPO_ALIASES.keys()),
        default=[],
        help="Import a single source (repeatable).",
    )
    ap.add_argument(
        "--overwrite",
        action="store_true",
        help="Replace existing Bronze rows for selected scope.",
    )
    args = ap.parse_args(argv)

    project_root = args.project_root or default_project_root or Path(__file__).resolve().parents[3]
    project_root = project_root.resolve()
    source_root = args.source_root.expanduser().resolve()

    if args.all:
        keys = list(REPO_ALIASES.keys())
    elif args.source:
        keys = list(args.source)
    else:
        print(
            "Specify --all or at least one --source {vanetik,dataturks,mehyar,nlp_ner}",
            file=sys.stderr,
        )
        sys.exit(2)

    try:
        summary = run_import(
            source_root=source_root,
            project_root=project_root,
            keys=keys,
            overwrite=args.overwrite,
        )
    except Exception as exc:
        logger.exception("Import failed: %s", exc)
        sys.exit(1)

    print(
        f"Bronze import done: {summary['resumes']} resumes, {summary['jobs']} jobs, "
        f"{summary['annotations']} NER rows. ground_truth rows written: {summary['ground_truth_rows']}"
    )


if __name__ == "__main__":
    run_cli()
